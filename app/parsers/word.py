from app.parsers.base import BaseParser, ParsedDocument, ParsedPage


class WordParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document as DocxDocument

        docx = DocxDocument(file_path)
        doc = ParsedDocument()
        current_content = []

        # 段落内容
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                current_content.append(text)

        # 表格内容
        for table in docx.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    current_content.append(row_text)

        full_text = self.clean_text("\n".join(current_content))
        if full_text:
            doc.pages.append(ParsedPage(page_num=1, content=full_text))
        return doc
